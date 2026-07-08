"""Convert the Vietnamese report (docs/report.md) to a .docx tuned for Google Docs.

Styling (all Times New Roman):
  Heading 1  14pt bold, UPPERCASE, centered   -> chapter titles and standalone sections
  Heading 2  14pt bold, left        (1.1. ...)
  Heading 3  13pt italic, left      (1.1.1. ...)
  Heading 4  13pt italic, centered  -> figure captions
  Heading 5  13pt italic, centered  -> table captions
  Normal     13pt, justified

Front matter after "Lời nói đầu": Mục lục (headings 1-3), Danh mục bảng (heading 5),
Danh mục hình ảnh (heading 4). A page-number footer starts at 1 on every page.

Mermaid blocks are rendered to images (via mermaid-cli / npx) when possible; otherwise a
placeholder is inserted so the diagram can be added by hand in Google Docs.

Run:
    uv run python scripts/md_to_docx.py                       # default in/out
    uv run python scripts/md_to_docx.py --no-render-mermaid   # skip rendering
"""

from __future__ import annotations

import argparse
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# --------------------------------------------------------------------------- config
FONT = "Times New Roman"
# Font for text inside mermaid diagrams. Times New Roman is used if present; otherwise
# Liberation Serif / DejaVu Serif (installed, Vietnamese-capable, TNR-metric-compatible).
DIAGRAM_FONT = "Times New Roman, Liberation Serif, DejaVu Serif, serif"
BLACK = RGBColor(0, 0, 0)
LINE_SPACING = 1.5           # line spacing for body paragraphs
IMAGE_WIDTH_IN = 5.7         # rendered diagram width
MAX_IMAGE_HEIGHT_IN = 6.5    # cap height so tall (portrait) diagrams shrink to fit a page
FIRST_LINE_INDENT_IN = 0.5   # first-line (tab) indent for body paragraphs

# (size_pt, bold, italic, alignment) per heading level.
HEADING = {
    1: (14, True, False, ALIGN.CENTER),
    2: (14, True, False, ALIGN.LEFT),
    3: (13, False, True, ALIGN.LEFT),
    4: (13, False, True, ALIGN.CENTER),   # figure captions
    5: (13, False, True, ALIGN.CENTER),   # table captions
}
BODY_SIZE = 13

DEFAULT_IN = Path("docs/report.md")
DEFAULT_OUT = Path("docs/report.docx")
IMAGES_DIR = Path("docs/images")
FONTS_DIR = Path("fonts")   # drop a Times New Roman .ttf here; no system install needed

CAPTION_RE = re.compile(r'<div align="center"><em>(.*?)</em></div>')
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


# --------------------------------------------------------------------------- helpers
def set_run_font(run, size=None, bold=None, italic=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def style_font(style, size, bold, italic):
    f = style.font
    f.name = FONT
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    f.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def build_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    style_font(normal, BODY_SIZE, False, False)
    pf = normal.paragraph_format
    pf.alignment = ALIGN.JUSTIFY
    pf.line_spacing = LINE_SPACING
    for level, (size, bold, italic, align) in HEADING.items():
        st = doc.styles[f"Heading {level}"]
        style_font(st, size, bold, italic)
        st.paragraph_format.alignment = align
        st.paragraph_format.line_spacing = LINE_SPACING


def parse_inline(text: str):
    """Yield (chunk, bold, italic) tuples from a line of inline markdown."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True, False
        elif part.startswith("*") and part.endswith("*"):
            yield part[1:-1], False, True
        elif part.startswith("`") and part.endswith("`"):
            yield part[1:-1], False, False
        else:
            yield part, False, False


def add_paragraph(doc, text, align=None, size=BODY_SIZE, italic=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for chunk, b, i in parse_inline(text):
        run = p.add_run(chunk)
        set_run_font(run, size=size, bold=b, italic=(italic if italic is not None else i))
    return p


def add_heading(doc, text, level, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    if level == 1:
        text = text.upper()
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_field(doc, instr):
    p = doc.add_paragraph()
    r = p.add_run()._r
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    r.append(begin); r.append(instr_el); r.append(sep)
    ph = p.add_run("(Cập nhật trường trong Word/LibreOffice để hiển thị)")
    set_run_font(ph, size=BODY_SIZE, italic=True)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    p.add_run()._r.append(end)
    return p


def set_update_fields(doc):
    settings = doc.settings.element
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    settings.append(el)


def add_footer_page_numbers(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    p = section.footer.paragraphs[0]
    p.alignment = ALIGN.CENTER
    r = p.add_run()._r
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r.append(begin); r.append(instr); r.append(end)
    set_run_font(p.runs[0], size=BODY_SIZE)
    pg = OxmlElement("w:pgNumType"); pg.set(qn("w:start"), "1")
    section._sectPr.append(pg)


# --------------------------------------------------------------------------- mermaid
def setup_fontconfig():
    """Let a .ttf dropped in fonts/ be used by the renderer, alongside system fonts.

    Nothing is installed on the machine; fontconfig is just pointed at fonts/ for this
    process (and the Chromium it launches).
    """
    if not FONTS_DIR.exists():
        return
    conf = Path(tempfile.gettempdir()) / "kmam_report_fonts.conf"
    conf.write_text(
        '<?xml version="1.0"?>\n'
        '<!DOCTYPE fontconfig SYSTEM "fonts.dtd">\n'
        "<fontconfig>\n"
        f"  <dir>{FONTS_DIR.resolve()}</dir>\n"
        '  <include ignore_missing="yes">/etc/fonts/fonts.conf</include>\n'
        "</fontconfig>\n",
        encoding="utf-8",
    )
    os.environ.setdefault("FONTCONFIG_FILE", str(conf))


def render_mermaid(code, out_png):
    """Render a mermaid block to PNG via mermaid-cli (Playwright/Chromium, no Node)."""
    try:
        from mermaid_cli import render_mermaid_file_sync
    except Exception:
        return False
    with tempfile.TemporaryDirectory() as tmp:
        mmd = Path(tmp) / "d.mmd"
        mmd.write_text(code, encoding="utf-8")
        try:
            render_mermaid_file_sync(
                str(mmd), str(out_png), "png",
                background_color="white",
                mermaid_config={"fontFamily": DIAGRAM_FONT,
                                "themeVariables": {"fontFamily": DIAGRAM_FONT}},
                playwright_config={"args": ["--no-sandbox"]},
            )
        except Exception as exc:
            print(f"  [mermaid] render failed: {type(exc).__name__}: {exc}")
            return False
    return out_png.exists()


def add_diagram(doc, code, index, render_enabled):
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    out = IMAGES_DIR / f"hinh-{index}.png"
    img = out if out.exists() else None
    if img is None and render_enabled and render_mermaid(code, out):
        img = out
    p = doc.add_paragraph()
    p.alignment = ALIGN.CENTER
    if img is not None:
        pic = p.add_run().add_picture(str(img), width=Inches(IMAGE_WIDTH_IN))
        # Cap the height so tall (portrait) diagrams shrink to fit within a page,
        # keeping the aspect ratio. Wide diagrams stay at the full width.
        max_h = Inches(MAX_IMAGE_HEIGHT_IN)
        if pic.height > max_h:
            pic.width = Emu(int(pic.width * max_h / pic.height))
            pic.height = max_h
        print(f"  [diagram {index}] embedded {img}")
    else:
        run = p.add_run(f"[Sơ đồ Hình {index} — chèn ảnh thủ công tại đây]")
        set_run_font(run, size=BODY_SIZE, italic=True)
        print(f"  [diagram {index}] placeholder (no image / render unavailable)")


# --------------------------------------------------------------------------- tables
def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = []
    for line in rows[2:]:
        body.append([c.strip() for c in line.strip().strip("|").split("|")])
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for j, cell_text in enumerate(header):
        _fill_cell(table.rows[0].cells[j], cell_text, bold=True)
    for r in body:
        cells = table.add_row().cells
        for j in range(len(header)):
            _fill_cell(cells[j], r[j] if j < len(r) else "", bold=False)


def _fill_cell(cell, text, bold):
    cell.text = ""
    p = cell.paragraphs[0]
    for chunk, b, i in parse_inline(text):
        run = p.add_run(chunk)
        set_run_font(run, size=BODY_SIZE, bold=(bold or b), italic=i)


# --------------------------------------------------------------------------- front matter
def emit_frontmatter(doc):
    specs = [
        ("MỤC LỤC", 'TOC \\o "1-3" \\h \\z \\u'),
        ("DANH MỤC BẢNG", 'TOC \\h \\z \\t "Heading 5,1"'),
        ("DANH MỤC HÌNH ẢNH", 'TOC \\h \\z \\t "Heading 4,1"'),
    ]
    for title, instr in specs:
        doc.add_page_break()
        add_heading(doc, title, 1)
        add_field(doc, instr)
    # No trailing page break: the first chapter (Heading 1) breaks onto a new page itself.


# --------------------------------------------------------------------------- main parse
def convert(md_path: Path, out_path: Path, render_enabled: bool):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    if render_enabled:
        setup_fontconfig()
    doc = Document()
    build_styles(doc)

    buffer: list[str] = []
    fig_index = 0
    right_align = False
    frontmatter_done = False
    first_h1_done = False

    def flush():
        nonlocal buffer
        if buffer:
            p = add_paragraph(doc, " ".join(buffer),
                              align=ALIGN.RIGHT if right_align else None)
            # First-line (tab) indent for ordinary body paragraphs, but not for the
            # right-aligned signature block.
            if not right_align:
                p.paragraph_format.first_line_indent = Inches(FIRST_LINE_INDENT_IN)
            buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Insert the front matter right before the first chapter.
        if stripped.startswith("# CHƯƠNG") and not frontmatter_done:
            flush()
            emit_frontmatter(doc)
            frontmatter_done = True

        if not stripped:
            flush()
            i += 1
            continue

        # Figure / table captions -> Heading 4 / Heading 5.
        m = CAPTION_RE.match(stripped)
        if m:
            flush()
            text = m.group(1).strip()
            add_heading(doc, text, 5 if text.startswith("Bảng") else 4)
            i += 1
            continue

        if stripped.startswith("<div align=\"right\""):
            flush(); right_align = True; i += 1; continue
        if stripped == "</div>":
            flush(); right_align = False; i += 1; continue

        # Mermaid code fence.
        if stripped.startswith("```mermaid"):
            flush()
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1  # skip closing fence
            fig_index += 1
            add_diagram(doc, "\n".join(code), fig_index, render_enabled)
            continue
        if stripped.startswith("```"):  # any other fence: skip fence markers
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buffer.append(lines[i]); i += 1
            i += 1
            flush()
            continue

        # Headings.
        if stripped.startswith("#"):
            flush()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            # A level-1 heading (chapter / standalone section) starts on a new page,
            # except the very first heading of the document (Lời nói đầu). The first
            # chapter already lands on a new page after the front matter.
            page_break = False
            if level == 1:
                page_break = first_h1_done
                first_h1_done = True
            add_heading(doc, text, min(level, 3), page_break=page_break)
            i += 1
            continue

        # Markdown pipe table.
        if stripped.startswith("|"):
            flush()
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue

        # Normal text (accumulate wrapped lines into one paragraph).
        buffer.append(stripped)
        i += 1

    flush()
    set_update_fields(doc)
    add_footer_page_numbers(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path} ({fig_index} diagrams processed)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=DEFAULT_IN)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--no-render-mermaid", action="store_true",
                        help="do not attempt to render mermaid diagrams")
    args = parser.parse_args()
    convert(args.input, args.output, render_enabled=not args.no_render_mermaid)


if __name__ == "__main__":
    main()
